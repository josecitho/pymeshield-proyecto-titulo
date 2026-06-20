import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

md_path = r"C:\Users\josev\Escritorio\PymeShield Beta\MANUAL_DE_USO.md"
docx_path = r"C:\Users\josev\Escritorio\PymeShield Beta\MANUAL_DE_USO.docx"

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, fill_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), fill_hex)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Initialize document
doc = docx.Document()

# Adjust margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Document Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("MANUAL DE USO Y REMEDIACIÓN — PYMESHIELD")
run.font.name = "Outfit"
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(10, 37, 64) # Navy blue

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
s_run = subtitle.add_run("Guía de Instalación, Configuración de Credenciales y Operación del Panel")
s_run.font.name = "Outfit"
s_run.font.size = Pt(11)
s_run.font.italic = True
s_run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph() # Spacing

# Credentials Highlight Box (Table style)
table = doc.add_table(rows=1, cols=1)
table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
cell = table.cell(0, 0)
set_cell_background(cell, "F2F5F8") # Soft light blue gray background
set_cell_margins(cell, top=140, bottom=140, left=200, right=200)

p_box = cell.paragraphs[0]
r_box = p_box.add_run("🔑 CREDENCIALES DE ACCESO DE FÁBRICA (PRIMER USO)")
r_box.font.name = "Outfit"
r_box.font.size = Pt(11)
r_box.font.bold = True
r_box.font.color.rgb = RGBColor(197, 90, 17) # Orange accent

p_user = cell.add_paragraph()
p_user.add_run("• Usuario por defecto: ").bold = True
p_user.add_run("admin")
p_pass = cell.add_paragraph()
p_pass.add_run("• Contraseña de fábrica: ").bold = True
p_pass.add_run("pymeshield2026")
p_mfa = cell.add_paragraph()
p_mfa.add_run("• Doble Factor (MFA): ").bold = True
p_mfa.add_run("Al ingresar por primera vez, el sistema mostrará un código QR. Escanéelo con la aplicación Google Authenticator en su celular para vincular su cuenta.")

p_box2 = cell.add_paragraph()
p_box2.paragraph_format.space_before = Pt(10)
r_box2 = p_box2.add_run("⚙️ CÓMO CAMBIAR LAS CREDENCIALES POR SEGURIDAD:")
r_box2.font.name = "Outfit"
r_box2.font.size = Pt(11)
r_box2.font.bold = True
r_box2.font.color.rgb = RGBColor(10, 37, 64)

p_step1 = cell.add_paragraph()
p_step1.add_run("1. Inicie sesión en PymeShield con las credenciales arriba indicadas.\n")
p_step1.add_run("2. En el menú lateral izquierdo, haga clic en la pestaña ")
r_ajustes = p_step1.add_run("Ajustes de Acceso")
r_ajustes.italic = True
p_step1.add_run(".\n")
p_step1.add_run("3. Rellene los campos: clave actual (")
r_key = p_step1.add_run("pymeshield2026")
r_key.font.name = "Courier New"
p_step1.add_run("), escriba su nueva contraseña y repítala en el campo de confirmación.\n")
p_step1.add_run("4. Presione el botón azul ")
r_btn = p_step1.add_run("Guardar Cambios")
r_btn.bold = True
p_step1.add_run(" para guardar el hash de forma segura en disco.")

doc.add_paragraph() # Spacing

# Read MANUAL_DE_USO.md content and parse headings
with open(md_path, "r", encoding="utf-8") as f:
    lines = f.readlines()

in_section_1 = False
for line in lines:
    line = line.strip()
    if not line:
        continue
    if line.startswith("# "):
        # Skip main title since we already added it
        continue
    elif line.startswith("## "):
        in_section_1 = ("1. Instalación" in line)
        h_text = line.replace("## ", "").replace("—", "-")
        p = doc.add_heading(level=1)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(h_text)
        r.font.name = "Outfit"
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(10, 37, 64)
    elif line.startswith("### "):
        h_text = line.replace("### ", "")
        p = doc.add_heading(level=2)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(h_text)
        r.font.name = "Outfit"
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = RGBColor(197, 90, 17)
    elif line.startswith("* ") or line.startswith("- "):
        p_text = line[2:]
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(p_text)
        r.font.name = "Outfit"
        r.font.size = Pt(10.5)
    elif line[0].isdigit() and line[1:3] == ". ":
        # Numbered list
        p_text = line[3:]
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(line[0:2] + " " + p_text)
        r.font.name = "Outfit"
        r.font.size = Pt(10.5)
    else:
        # Standard paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(line)
        r.font.name = "Outfit"
        r.font.size = Pt(10.5)

# Save docx
doc.save(docx_path)
print("MANUAL_DE_USO.docx generated successfully!")

# Remove original markdown manual to keep beta directory clean
os.remove(md_path)
print("Removed temporary MANUAL_DE_USO.md from PymeShield Beta.")
