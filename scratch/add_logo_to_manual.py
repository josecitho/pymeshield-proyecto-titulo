import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

docx_path = r"C:\Users\josev\Escritorio\PymeShield Beta\MANUAL_DE_USO.docx"
logo_path = r"C:\Users\josev\Escritorio\PymeShield Beta\public\logo.png"

# Since we already deleted MANUAL_DE_USO.md, we will read the text from the current DOCX 
# or rewrite the generation by using a static string of the markdown or extracting from docx.
# Wait, let's just extract the paragraphs from the current docx, recreate a new one, 
# insert the logo at the very beginning, and save it! This is very easy and doesn't require MD.

print("Opening existing docx...")
old_doc = docx.Document(docx_path)

new_doc = docx.Document()

# Adjust margins
for section in new_doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# 1. Add Logo centered
if os.path.exists(logo_path):
    print("Embedding logo image at the top...")
    p_logo = new_doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(10)
    p_logo.paragraph_format.space_after = Pt(12)
    run_logo = p_logo.add_run()
    run_logo.add_picture(logo_path, width=Inches(1.2)) # Centered 1.2 inches wide logo
else:
    print("Logo image not found!")

# 2. Copy the rest of paragraphs and tables from old doc
# The first paragraph of the old doc is the title: "MANUAL DE USO Y REMEDIACIÓN — PYMESHIELD"
# The second is the subtitle.
# The third is a spacing.
# Then the credentials table.
# Then the rest of paragraphs.

# Let's copy the elements in order
# First, copy the table (credentials box) by recreating it
print("Recreating credentials table...")
old_table = old_doc.tables[0]
new_table = new_doc.add_table(rows=1, cols=1)
new_table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER

# Helper functions for cell style
def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, fill_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), fill_hex)
    cell._element.get_or_add_tcPr().append(shading_elm)

set_cell_background(new_table.cell(0,0), "F2F5F8")
set_cell_margins(new_table.cell(0,0), top=140, bottom=140, left=200, right=200)

# Copy table paragraphs
cell_p_list = old_table.cell(0,0).paragraphs
new_cell = new_table.cell(0,0)
for idx, p in enumerate(cell_p_list):
    if idx == 0:
        new_p = new_cell.paragraphs[0]
    else:
        new_p = new_cell.add_paragraph()
    new_p.paragraph_format.space_before = p.paragraph_format.space_before
    new_p.paragraph_format.space_after = p.paragraph_format.space_after
    for r in p.runs:
        new_run = new_p.add_run(r.text)
        new_run.bold = r.bold
        new_run.italic = r.italic
        new_run.font.name = r.font.name
        new_run.font.size = r.font.size
        new_run.font.color.rgb = r.font.color.rgb

# Copy other paragraphs
print("Copying remaining text paragraphs...")
for idx, p in enumerate(old_doc.paragraphs):
    # Copy paragraph properties
    new_p = new_doc.add_paragraph()
    new_p.alignment = p.alignment
    new_p.paragraph_format.space_before = p.paragraph_format.space_before
    new_p.paragraph_format.space_after = p.paragraph_format.space_after
    new_p.paragraph_format.left_indent = p.paragraph_format.left_indent
    new_p.style = p.style
    
    for r in p.runs:
        new_run = new_p.add_run(r.text)
        new_run.bold = r.bold
        new_run.italic = r.italic
        new_run.font.name = r.font.name
        new_run.font.size = r.font.size
        new_run.font.color.rgb = r.font.color.rgb

# Save the new document
new_doc.save(docx_path)
print("Manual generated with embedded logo successfully!")
