import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

docx_path_dev = r"C:\Users\josev\Escritorio\ProyectoPymeShield01\ProyectoPymeShield\GUIA_TECNICA_DE_DESARROLLO.docx"

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, fill_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), fill_hex)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_heading_styled(doc, text, level):
    p = doc.add_heading(level=level)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Outfit"
    if level == 1:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(10, 37, 64) # Navy blue
    else:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(197, 90, 17) # Orange
    return p

def add_paragraph_styled(doc, text, bold_prefix="", bullet=False, code=False):
    style = 'List Bullet' if bullet else 'Normal'
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Outfit"
        r_pre.font.size = Pt(10.5)
        r_pre.bold = True
        r_pre.font.color.rgb = RGBColor(10, 37, 64)
        
    r = p.add_run(text)
    if code:
        r.font.name = "Courier New"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(150, 40, 40)
    else:
        r.font.name = "Outfit"
        r.font.size = Pt(10.5)
    return p

def build_docx(path):
    doc = docx.Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("GUÍA DE ESTUDIO TÉCNICO Y ARQUITECTURA\nPYMESHIELD")
    tr.font.name = "Outfit"
    tr.font.size = Pt(20)
    tr.font.bold = True
    tr.font.color.rgb = RGBColor(10, 37, 64)
    
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    str_run = st.add_run("Bitácora de Ingeniería de Software, Flujos de Ciberseguridad y Cuestionario de Defensa Académica")
    str_run.font.name = "Outfit"
    str_run.font.size = Pt(11)
    str_run.font.italic = True
    str_run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph() # Spacing
    
    # ----------------------------------------------------
    add_heading_styled(doc, "Sección 1: Arquitectura Desacoplada y APIs de Red", level=1)
    add_paragraph_styled(doc, "PymeShield está diseñado bajo un modelo de arquitectura web moderna de dos capas principales (Backend y Frontend desacoplados) que se comunican mediante servicios API REST e intercambio de datos estructurados en formato JSON.")
    
    add_paragraph_styled(doc, "El Frontend (Cliente) está construido con HTML5 semántico, CSS3 de alto impacto (Modo NOC, animaciones fluidas) y JavaScript nativo. Su función exclusiva es pintar la interfaz en el navegador y capturar las interacciones del administrador.", bold_prefix="• Frontend (HTML/CSS/JS): ")
    add_paragraph_styled(doc, "El Backend (Servidor) está construido sobre Node.js y Express. Se encarga de conectarse con el sistema operativo Windows, ejecutar los sockets UDP/ARP de escaneo, alterar las reglas del Firewall del sistema y persistir los datos en la base de datos SQLite usando Prisma ORM.", bold_prefix="• Backend (Node.js/Express): ")
    
    add_heading_styled(doc, "¿Por qué son necesarias las APIs? (Justificación Técnica)", level=2)
    add_paragraph_styled(doc, "Por razones de seguridad del navegador, las aplicaciones web corren en un sandbox (caja de arena) aislado. Un navegador web no tiene permisos de sistema para enviar paquetes UDP sin procesar, consultar la tabla de caché física ARP de la máquina o manipular el registro de Windows local. Por ende, la interfaz requiere de las APIs como un puente de comunicación para indicarle al servidor Node.js que realice estas acciones en su nombre y le devuelva la información consolidada.", bullet=True)
    add_paragraph_styled(doc, "Al tener funciones separadas por APIs (como /api/scan o /api/devices), el software es altamente escalable. En el futuro, el backend de PymeShield no tendrá que ser modificado para construir una aplicación móvil en Google Play Store; la app móvil simplemente consumirá estos mismos endpoints REST para mostrar alertas en el teléfono del usuario.", bullet=True)
    
    add_heading_styled(doc, "Catálogo Completo de Endpoints API de PymeShield", level=2)
    add_paragraph_styled(doc, "A continuación, se detalla el catálogo de APIs programadas en el servidor server.js que interactúan con el panel:")
    
    add_paragraph_styled(doc, "Obtiene la lista de todos los dispositivos de la subred local.", bold_prefix="• GET /api/devices: ", bullet=True)
    add_paragraph_styled(doc, "Registra o revoca un dispositivo de la lista de confianza, asignando su alias correspondiente.", bold_prefix="• POST /api/devices/toggle-authorize: ", bullet=True)
    add_paragraph_styled(doc, "Aplica o remueve un bloqueo de red a nivel de firewall para contención de intrusos.", bold_prefix="• POST /api/devices/block: ", bullet=True)
    add_paragraph_styled(doc, "Genera y descarga un script dinámico de remediación (.bat/.sh) para el host seleccionado.", bold_prefix="• GET /api/devices/:id/hardening-script: ", bullet=True)
    add_paragraph_styled(doc, "Inicia el motor de escaneo de red (barrido UDP + lectura de caché ARP).", bold_prefix="• POST /api/scan: ", bullet=True)
    add_paragraph_styled(doc, "Obtiene el historial de auditorías de red previas para el gráfico de tendencias.", bold_prefix="• GET /api/scans/history: ", bullet=True)
    add_paragraph_styled(doc, "Inyecta un dispositivo atacante simulado para la demostración práctica de incidentes.", bold_prefix="• POST /api/scan/simulate-attack: ", bullet=True)
    add_paragraph_styled(doc, "Obtiene la configuración activa del sistema (Demo Mode, Zero-Trust, Webhook).", bold_prefix="• GET /api/settings: ", bullet=True)
    add_paragraph_styled(doc, "Actualiza las directivas del sistema (Zero-Trust y URL de Webhook).", bold_prefix="• POST /api/settings/update: ", bullet=True)
    add_paragraph_styled(doc, "Activa o desactiva de forma dinámica el modo demostración académica.", bold_prefix="• POST /api/settings/toggle-demo: ", bullet=True)
    add_paragraph_styled(doc, "Valida claves de licenciamiento del software para activar el modo Premium.", bold_prefix="• POST /api/settings/activate-license: ", bullet=True)
    add_paragraph_styled(doc, "Modifica de forma segura la contraseña del administrador aplicando hash SHA-256.", bold_prefix="• POST /api/settings/change-password: ", bullet=True)
    add_paragraph_styled(doc, "Restablece y desvincula el dispositivo móvil del Doble Factor (MFA).", bold_prefix="• POST /api/settings/reset-mfa: ", bullet=True)
    add_paragraph_styled(doc, "Envía un JSON payload de prueba para validar la integración con el webhook del SOAR.", bold_prefix="• POST /api/settings/test-webhook: ", bullet=True)
    add_paragraph_styled(doc, "Valida las credenciales administrativas de acceso iniciales en el login.", bold_prefix="• POST /api/login: ", bullet=True)
    add_paragraph_styled(doc, "Genera la semilla criptográfica Base32 y el código QR de vinculación del MFA.", bold_prefix="• POST /api/login/mfa-setup: ", bullet=True)
    add_paragraph_styled(doc, "Valida el código dinámico de 6 dígitos del autenticador para otorgar acceso.", bold_prefix="• POST /api/login/mfa: ", bullet=True)
    add_paragraph_styled(doc, "Obtiene el registro de auditoría técnica simplificada en español (Audit Log).", bold_prefix="• GET /api/audit-logs: ", bullet=True)
    add_paragraph_styled(doc, "Obtiene la lista de alertas de red activas.", bold_prefix="• GET /api/alerts: ", bullet=True)
    add_paragraph_styled(doc, "Marca todas las alertas críticas actuales como leídas.", bold_prefix="• POST /api/alerts/read-all: ", bullet=True)
    add_paragraph_styled(doc, "Descarga el reporte de cumplimiento PDF formal (Ley N° 21.719 chilena).", bold_prefix="• GET /api/reports/pdf: ", bullet=True)

    # ----------------------------------------------------
    add_heading_styled(doc, "Sección 2: Autoprotección y Capas de Seguridad Internas", level=1)
    add_paragraph_styled(doc, "Dado que PymeShield es una herramienta de auditoría de seguridad, se dotó al software de estrictos controles internos para evitar que sea manipulado o explotado por atacantes locales:")
    
    add_paragraph_styled(doc, "PymeShield nunca almacena contraseñas en texto plano. La clave de administración se almacena como un hash criptográfico SHA-256 no reversible en credenciales.json. Durante el inicio de sesión, el backend aplica la misma función hash al texto ingresado y compara el resultado, garantizando que el secreto no pueda ser recuperado del disco.", bold_prefix="1. Hashing de Claves (SHA-256): ")
    add_paragraph_styled(doc, "Como segundo factor criptográfico, implementamos el algoritmo TOTP (Time-Based One-Time Password) mediante la librería otplib. En el primer uso, el servidor genera una clave secreta Base32 y la codifica en un código QR dinámico. El administrador lo vincula con Google Authenticator. En ingresos posteriores, el servidor requiere la validación del código dinámico de 6 dígitos basado en la hora matemática.", bold_prefix="2. Doble Factor de Autenticación Real (MFA): ")
    add_paragraph_styled(doc, "Al ejecutarse localmente y ser empaquetada en una red física cerrada sin dependencias de la nube, la superficie de ataque externa se reduce al mínimo. Para mitigar ataques locales de red, las APIs escuchan de forma restringida y la base de datos SQLite local (dev.db) requiere permisos a nivel de sistema operativo.", bold_prefix="3. Aislamiento Local e Invariabilidad: ")
    
    # ----------------------------------------------------
    add_heading_styled(doc, "Sección 3: Estructura del Directorio del Proyecto (Mapa de Carpetas)", level=1)
    add_paragraph_styled(doc, "Es vital entender qué contiene cada carpeta del proyecto PymeShield de cara a la defensa del software:")
    
    add_paragraph_styled(doc, "Contiene todas las dependencias y librerías externas de Node.js instaladas por npm (como Express para el servidor, Prisma para la base de datos, Otplib para TOTP, QRCode para los códigos QR, PDFKit para generar reportes y WS para WebSockets de alertas).", bold_prefix="• node_modules/: ", bullet=True)
    add_paragraph_styled(doc, "Contiene el motor de persistencia. Incluye schema.prisma (donde se definen las tablas de base de datos Device, Port, Alert, AuditLog y Recommendation) y dev.db (el archivo físico de base de datos SQLite local donde se guarda toda la información sincronizada).", bold_prefix="• prisma/: ", bullet=True)
    add_paragraph_styled(doc, "Contiene todos los archivos del cliente (Frontend). Incluye index.html (la estructura del panel), style.css (los estilos visuales, Modo NOC y animaciones), app.js (el cerebro del navegador que dibuja el mapa SVG y procesa las llamadas de APIs) y logo.png (el logotipo corporativo).", bold_prefix="• public/: ", bullet=True)
    add_paragraph_styled(doc, "El punto de entrada principal del Backend de PymeShield. Inicia el servidor, define las rutas de las APIs, ejecuta el escaneo de red por UDP y realiza las peticiones a la base de datos SQLite.", bold_prefix="• server.js: ", bullet=True)
    add_paragraph_styled(doc, "Almacena los ajustes administrativos del sistema (el hash SHA-256 de la contraseña, el secreto Base32 de la MFA, la URL del webhook SOAR y el estado de la directiva Zero-Trust).", bold_prefix="• credenciales.json: ", bullet=True)
    add_paragraph_styled(doc, "Archivos de configuración de Node.js donde se listan las dependencias del proyecto, autoría, scripts de inicio y versiones bloqueadas.", bold_prefix="• package.json y package-lock.json: ", bullet=True)
    add_paragraph_styled(doc, "El instalador offline oficial de Node.js para Windows x64 de 64 bits.", bold_prefix="• node-v20.11.0-x64.msi: ", bullet=True)
    add_paragraph_styled(doc, "Ícono en alta resolución de PymeShield utilizado para asociarlo de forma automática al acceso directo del Escritorio.", bold_prefix="• pymeshield.ico: ", bullet=True)
    add_paragraph_styled(doc, "Scripts de un solo clic para instalar Node.js offline, levantar dependencias y base de datos (Instalar), e iniciar el servidor limpiando el puerto 3000 y abriendo el navegador (Iniciar).", bold_prefix="• Instalar_PymeShield.bat e Iniciar PymeShield.bat: ", bullet=True)

    # ----------------------------------------------------
    add_heading_styled(doc, "Sección 4: El Viaje del Código (Flujo Paso a Paso de la Aplicación)", level=1)
    add_paragraph_styled(doc, "A continuación, se detalla cronológicamente la historia de ejecución del software desde el primer clic del usuario:")
    
    add_paragraph_styled(doc, "El usuario conecta el pendrive, copia la carpeta PymeShield Beta localmente y abre Instalar_PymeShield.bat. El script utiliza consultas de consola (where node) para verificar requisitos. Si no está instalado, ejecuta de forma silenciosa el archivo node-v20.11.0-x64.msi local mediante msiexec.exe de Windows, instala las dependencias de node_modules, sincroniza las tablas de la base de datos con prisma db push, y genera el acceso directo en el Escritorio usando la interfaz COM de WScript.Shell de Windows.", bold_prefix="Fase 1: Instalación Autónoma y Offline: ")
    add_paragraph_styled(doc, "El usuario abre la aplicación y digita admin / pymeshield2026. El servidor valida la clave actual, genera el secreto MFA y requiere el escaneo QR. Una vez validado, el estado se guarda permanentemente en credenciales.json. En Ajustes, el administrador puede cambiar la clave, lo que regenera el hash SHA-256 de forma interactiva en caliente.", bold_prefix="Fase 2: Autenticación y Cambio de Claves: ")
    add_paragraph_styled(doc, "El usuario hace clic en 'Escanear Red'. El servidor Node.js (server.js) utiliza el módulo dgram para enviar en paralelo un paquete UDP vacío de 1 byte a la IP de descarte (puerto 9) de todos los hosts posibles en la subred (ej: 192.168.1.1 a 192.168.1.254). Esto obliga a las tarjetas de red de los hosts a responder resolviendo la dirección MAC, actualizando la caché ARP del sistema operativo Windows local. Luego, el backend lee esta caché y entrega una lista precisa de dispositivos activos, evadiendo los bloqueos de firewall de ping (ICMP).", bold_prefix="Fase 3: El Escaneo Avanzado por UDP/ARP: ")
    add_paragraph_styled(doc, "En el mapa topológico interactivo en SVG, los dispositivos descubiertos orbitan en una elipse animada al router. El sistema compara sus MACs contra el inventario SQLite. Si hay dispositivos no autorizados (isAuthorized = false), el mapa los pinta de rojo brillante con una línea de conexión discontinua (dashed) y se gatilla un banner de advertencia rojo arriba en el dashboard.", bold_prefix="Fase 4: Lista Blanca y Alertas de Intrusos: ")
    add_paragraph_styled(doc, "El usuario activa el switch de control Zero-Trust (NAC). Al detectar un intruso en el escaneo, PymeShield ejecuta automáticamente comandos netsh para bloquear esa IP en el firewall, envía un JSON con la alerta vía Webhook a su central SOAR/Slack externa y genera un log de auditoría. Para mitigar brechas, el usuario puede descargar un script de Hardening (.bat o .sh) generado dinámicamente según los puertos abiertos del equipo para cerrarlos con un clic.", bold_prefix="Fase 5: Contención Zero-Trust (NAC) y Mitigación: ")
    add_paragraph_styled(doc, "El usuario activa el Modo NOC. El estilo visual pasa a negro y rojo neón, y la app utiliza la Web Audio API del navegador para sintetizar ondas sonoras de sirena a nivel de hardware. Además, permite descargar un Reporte PDF formal con firma criptográfica para evidenciar el cumplimiento legal ante fiscalizaciones de ciberseguridad.", bold_prefix="Fase 6: Bitácora de Emergencias y Reportabilidad: ")
    
    # ----------------------------------------------------
    add_heading_styled(doc, "Sección 5: Cuestionario de Defensa (Preguntas de la Comisión)", level=1)
    
    p1 = doc.add_paragraph()
    p1.add_run("Pregunta 1: ¿Por qué su escáner detecta dispositivos que tienen el Firewall de Windows activado, mientras que herramientas tradicionales fallan?").bold = True
    p1_ans = doc.add_paragraph()
    p1_ans.add_run("Respuesta: ").italic = True
    p1_ans.add_run("Las herramientas tradicionales envían paquetes ICMP Echo Request (pings). Por defecto, el Firewall de Windows y los dispositivos móviles descartan estos paquetes. PymeShield soluciona esto enviando solicitudes de conexión UDP en el puerto 9 de forma paralela. Aunque el firewall descarte el puerto, la tarjeta de red se ve forzada a resolver su dirección física a nivel de Capa 2 (ARP) para procesar el paquete. Esto actualiza la tabla caché ARP del servidor, permitiendo una identificación exacta de todos los hosts conectados.")
    
    p2 = doc.add_paragraph()
    p2.add_run("Pregunta 2: ¿Cómo funciona el aislamiento de red de PymeShield? ¿Realiza envenenamiento ARP (ARP Spoofing)?").bold = True
    p2_ans = doc.add_paragraph()
    p2_ans.add_run("Respuesta: ").italic = True
    p2_ans.add_run("No, PymeShield no realiza envenenamiento ARP ya que esa es una técnica ofensiva de hackeo que puede desestabilizar la red de producción del CESFAM o el colegio. En su lugar, el aislamiento se realiza de forma limpia mediante reglas defensivas de cortafuegos local (Host-Based NAC) inyectando reglas de bloqueo en el Firewall de Windows mediante comandos de sistema netsh, impidiendo cualquier comunicación entrante o saliente del host sospechoso hacia el panel o la puerta de enlace.")
    
    p3 = doc.add_paragraph()
    p3.add_run("Pregunta 3: ¿Qué utilidad tiene el Double Factor (MFA) si la aplicación es de uso local?").bold = True
    p3_ans = doc.add_paragraph()
    p3_ans.add_run("Respuesta: ")
    p3_ans.add_run("Aunque sea de ejecución local, la computadora que aloja el servidor PymeShield suele ser compartida o estar expuesta al paso del personal administrativo. Si un atacante interno logra robar el archivo credenciales.json o adivinar la clave de acceso, el Doble Factor basado en TOTP asegura que no podrá alterar las políticas de red ni desactivar la seguridad sin el teléfono celular físico del administrador.")
    
    p4 = doc.add_paragraph()
    p4.add_run("Pregunta 4: ¿Por qué decidieron estructurar el manual de instalación en un archivo de Word (.docx) en lugar de PDF o Markdown?").bold = True
    p4_ans = doc.add_paragraph()
    p4_ans.add_run("Respuesta: ")
    p4_ans.add_run("El objetivo de PymeShield es la usabilidad en organizaciones sin departamentos de IT avanzados. El formato Markdown (.md) requiere herramientas de renderizado específicas que confunden al usuario común abriéndose en texto plano en el Bloc de Notas. El archivo Word (.docx) es un formato corporativo nativo que cualquier usuario sabe abrir en Microsoft Word, permitiendo incluir un diseño limpio con el logo oficial del proyecto y las credenciales por defecto destacadas al inicio de la primera página de forma sumamente accesible.")
    
    p5 = doc.add_paragraph()
    p5.add_run("Pregunta 5: ¿Es su aplicación multiplataforma o solo funciona en Windows?").bold = True
    p5_ans = doc.add_paragraph()
    p5_ans.add_run("Respuesta: ")
    p5_ans.add_run("El núcleo web de PymeShield (Node.js, Express y SQLite) es 100% multiplataforma y puede correr nativamente en Windows, Linux y macOS, incluso empaquetado en contenedores Docker. Sin embargo, los módulos que interactúan directamente con el hardware y el sistema operativo (como la lectura de la caché ARP, el bloqueo de firewall de intrusos y los scripts de hardening) detectan dinámicamente el sistema operativo del host para utilizar los comandos y herramientas de red nativas correspondientes (como netsh en Windows, iptables o ufw en Linux, y pfctl en macOS).")

    p6 = doc.add_paragraph()
    p6.add_run("Pregunta 6: ¿Cómo está estructurado el directorio de su software PymeShield y qué función cumple cada carpeta principal?").bold = True
    p6_ans = doc.add_paragraph()
    p6_ans.add_run("Respuesta: ")
    p6_ans.add_run("El directorio está estructurado en 3 subcarpetas y archivos raíz principales. La carpeta node_modules/ contiene las dependencias externas (servidor web Express, encriptación Otplib y WebSockets). La carpeta prisma/ es la capa de persistencia conteniendo schema.prisma (modelo de base de datos) y dev.db (la base de datos SQLite local). La carpeta public/ es el Frontend conteniendo el HTML, el CSS del diseño/Modo NOC y el JavaScript que dibuja la topología en SVG. En la raíz, server.js controla el Backend (sockets, firewall y base de datos) y credenciales.json almacena la configuración de seguridad local.")

    # Save docx
    doc.save(path)
    print(f"Guia saved to {path}")

# Run compilation for dev only
build_docx(docx_path_dev)
print("Study guide updated with folder structure mapping successfully!")
